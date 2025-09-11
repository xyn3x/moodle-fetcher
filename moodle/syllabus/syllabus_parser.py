from google import genai
import pdfplumber
import os
import time
import requests
import json
from docx import Document
from io import BytesIO
from bs4 import BeautifulSoup
from dotenv import load_dotenv, dotenv_values
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException

load_dotenv()

def table_to_string(table):
    if not table:
        return ""

    columns = []
    max_width = 0
    for col in zip(*table):
        for cell in col:
            if not cell:
                continue
            max_width = len(str(cell))
        columns.append(max_width)
    
    formatted_table = ""
    for row_pos, row in enumerate(table):
        for col_pos, pos in enumerate(row):
            if cell:
                cur_text = str(cell)
            else:
                cur_text = ""
            cur_text = cur_text.ljust(columns[col_pos])
            if col_pos != 0:
                formatted_table += "|"
            formatted_table += cur_text
        formatted_table += "\n"
        if row_pos == 0:
            for width in columns:
                formatted_table += "-" * width
                formatted_table += "|"
    return formatted_table

def extract_from_pdf(driver, pdf_url):
    cookies = {c['name']: c['value'] for c in driver.get_cookies()}
    res = requests.get(pdf_url, cookies=cookies)
    pdf_file = BytesIO(res.content)
    extracted_text = ''
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for cur_table in tables:
                converted_table = table_to_string(cur_table)
                extracted_text += converted_table + '\n'
            if page.extract_text():
                extracted_text += page.extract_text() + '\n'
    return extracted_text

def extract_from_docx(driver, docx_url):
    cookies = {c['name']: c['value'] for c in driver.get_cookies()}
    res = requests.get(docx_url, cookies=cookies)
    docs_file = Document(BytesIO(res.content))
    extracted_text = ''
    for pars in docs_file.paragraphs:
        if pars.text.strip():
            extracted_text += pars.text + "\n"
    for table in docs_file.tables:
        extracted_text += table_to_string(table) + "\n"
    return extracted_text

def extract_from_gdocs(driver, gdocs_url):
    cookies = {c['name']: c['value'] for c in driver.get_cookies()}
    res = requests.get(gdocs_url, cookies=cookies)
    soup = BeautifulSoup(res.text, 'html.parser')
    for not_text in soup(["style", "script"]):
        not_text.extract()
    soup_text = soup.get_text()
    lines = []
    for line in soup_text.splitlines():
        lines.append(line.strip())
    extracted_text = ""
    for line in lines:
        for cur_text in line.split(" "):
            if cur_text.strip():
                extracted_text += cur_text.strip() + " "
        extracted_text += "\n"
    return extracted_text

def ask(content):
    client = genai.Client(api_key=os.getenv("genai_api"))
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=f"Give me the answer as a json only. You need to analyze this text and give information of every assigment that affects the grade as a json file. Give information about each assignment: name (make it as a key in json), its weight (out of 100), date (if none or TBA, then None. only date YYYY-MM-DD, if YYYY is not defined then it is from the current academic year). For each assigment count its amount during the course from the table of course outline or any table that contain course schedule or description if it's homework. If there are more than 1 date for each assigment, then split them so that one assigments have one date. Also add info about attendance (as a assigment also), also in json (maximum absence, and penalty. give info about penalty only as a one number that is final grade dedicated percentage. if it's says that penalty is 1 grade, then it's 5%, otherwise give precise percent). Here the texts itself:\n {content}",
    )
    return response.text

def parse_syllabus(driver, url, course_linklist):
    syllabus_json = {}
    for course_link in course_linklist: 
        driver.get(course_link)
        time.sleep(2)
        course_name_container = driver.find_element(By.CLASS_NAME, "page-header-headings")
        course_name = course_name_container.find_element(By.XPATH, "./*[1]").get_attribute("innerHTML")
        linkcontainer_list = driver.find_elements(By.CLASS_NAME, "activityname")
        #print(linkcontainer_list)
        for cur_container in linkcontainer_list:
            try:
                cur_link = cur_container.find_element(By.TAG_NAME, "a")
            except:
                continue
            #print(cur_link)
            try: 
                cur_linkname = cur_link.find_element(By.TAG_NAME, "span")
            except:
                continue
            cur_linkname = (cur_linkname.get_attribute("innerHTML")).lower()
            if "syllabus" in cur_linkname or ("course" in cur_linkname and "manual" in cur_linkname):
                try:
                    url = cur_link.get_attribute("href")
                    if "docs" in url:
                        content = extract_from_gdocs(driver, url)
                    else:
                        try:
                            content = extract_from_pdf(driver, url)
                        except:
                            content = extract_from_docx(driver, url)
                    response = ask(content)
                    response = response[7:-3].replace('\n', '')
                    response = json.loads(response)
                    syllabus_json.update({course_name : response})
                except: 
                    driver.get(cur_link.get_attribute("href"))
                    try:
                        link_list = driver.find_elements(By.TAG_NAME, "a")
                    except:
                        continue
                    for cur_link in link_list:
                        cur_linkname = (cur_link.get_attribute("innerHTML")).lower()
                        if "syllabus" in cur_linkname or ("course" in cur_linkname and "manual" in cur_linkname):
                            url = cur_link.get_attribute("href")
                            if "docs" in url:
                                content = extract_from_gdocs(driver, url)
                            else:
                                try:
                                    content = extract_from_pdf(driver, url)
                                except Exception:
                                    content = extract_from_docx(driver, url)
                            response = ask(content)
                            response = response[7:-3].replace('\n', '')
                            response = json.loads(response)
                            syllabus_json.update({course_name : response})
    # removing duplicates 
    modified_syllabus_json = {}
    for key in syllabus_json.keys():
        cur_name = ""
        cmpname = 0
        component_name = ""
        for cur in key:
            if cur == ",":
                break
            if cur == "-":
                if cmpname:
                    break
                cmpname = 1
                continue
            if not cmpname: 
                cur_name += cur
            else:
                component_name += cur
        component_name = component_name.lower()
        if "lab" in component_name:
            continue
        if cur_name in modified_syllabus_json.keys():
            modified_syllabus_json[cur_name].update(syllabus_json[key])
        else:
            modified_syllabus_json.update({cur_name : syllabus_json[key]})
    syllabus_json = json.dumps(modified_syllabus_json, indent=4)
    return syllabus_json


